import openai
import logging
import re
from django.conf import settings
from ..models import StudyMaterial, Question, Answer
from PyPDF2 import PdfReader
from docx import Document
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import requests.exceptions
import os
import math

# Set up logging
logger = logging.getLogger(__name__)

# Set OpenAI API key
openai.api_key = getattr(settings, 'OPENAI_API_KEY', None)

def extract_text_from_file(file_path, file_type):
    """
    Extract text from various file types with better handling.
    """
    try:
        logger.info(f"Extracting text from file: {file_path}, type: {file_type}")
        
        if file_type == 'PDF':
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ''
                for page_num, page in enumerate(reader.pages):
                    try:
                        extracted = page.extract_text()
                        if extracted:
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += extracted + '\n\n'
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue
                
                logger.info(f"Extracted {len(text)} characters from PDF")
                return text
        elif file_type == 'DOCX':
            doc = Document(file_path)
            # Extract text from paragraphs
            paragraph_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Extract text from tables
            table_text = ''
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_text += row_text + '\n'
            
            combined_text = paragraph_text + '\n\n' + table_text
            logger.info(f"Extracted {len(combined_text)} characters from DOCX")
            return combined_text
        elif file_type == 'TXT':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                text = file.read()
                logger.info(f"Extracted {len(text)} characters from TXT")
                return text
        elif file_type in ['JPG', 'JPEG', 'PNG']:
            # For image files, we could implement OCR here if needed
            # Using a placeholder message for now
            logger.warning(f"Image file type {file_type} not supported for text extraction")
            return "This is an image file. Text extraction from images would require OCR implementation."
        else:
            logger.error(f"Unsupported file type: {file_type}")
            return None
    except Exception as e:
        logger.error(f"Error extracting text from file: {e}")
        return None

def chunk_text(text, max_chunk_size=3000, overlap=200):
    """
    Split text into overlapping chunks for better processing.
    """
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + max_chunk_size
        
        # Try to break at sentence boundaries
        if end < len(text):
            # Look for sentence endings near the chunk boundary
            for i in range(end, max(start + max_chunk_size // 2, end - 200), -1):
                if text[i] in '.!?':
                    end = i + 1
                    break
        
        chunk = text[start:end]
        chunks.append(chunk)
        
        # Move start with overlap
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks

def analyze_content_complexity(text):
    """
    Analyze text complexity to determine appropriate question difficulty.
    """
    # Simple metrics for complexity
    sentences = len(re.findall(r'[.!?]+', text))
    words = len(text.split())
    avg_sentence_length = words / max(sentences, 1)
    
    # Technical terms detection
    technical_patterns = [
        r'\b[A-Z]{2,}\b',  # Acronyms
        r'\b\w+\(\w+\)',   # Function calls or formulas
        r'\b\d+\.\d+\b',   # Decimal numbers
        r'\b[a-zA-Z]+\d+\b',  # Mixed alphanumeric
    ]
    
    technical_score = sum(len(re.findall(pattern, text)) for pattern in technical_patterns)
    
    if avg_sentence_length > 20 or technical_score > 10:
        return 'HARD'
    elif avg_sentence_length > 15 or technical_score > 5:
        return 'MEDIUM'
    else:
        return 'EASY'

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=4, max=10),
    retry=retry_if_exception_type(requests.exceptions.RequestException),
    reraise=True
)
def call_openai_api(messages, max_tokens=2000):
    """
    Call OpenAI API with retry logic.
    """
    if not openai.api_key:
        raise ValueError("OpenAI API key is not configured")
    
    logger.info("Making OpenAI API call...")
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=messages,
            temperature=0.7,
            max_tokens=max_tokens,
        )
        logger.info("OpenAI API call successful")
        return response
    except Exception as e:
        logger.error(f"OpenAI API call failed: {e}")
        raise

def generate_diverse_questions(material_title, subject="general", difficulty="MEDIUM"):
    """
    Generate diverse sample questions when AI service fails.
    """
    logger.info(f"Generating diverse sample questions for {material_title}")
    
    question_templates = {
        'conceptual': [
            f"What is the primary concept discussed in {material_title}?",
            f"How does the main theory in {material_title} apply to real-world situations?",
            f"What are the key principles outlined in {material_title}?",
        ],
        'analytical': [
            f"Compare and contrast the different approaches mentioned in {material_title}.",
            f"What are the advantages and disadvantages of the methods described in {material_title}?",
            f"How do the concepts in {material_title} relate to each other?",
        ],
        'application': [
            f"In what scenarios would you apply the knowledge from {material_title}?",
            f"What practical problems can be solved using the information in {material_title}?",
            f"How would you implement the strategies discussed in {material_title}?",
        ]
    }
    
    sample_questions = []
    question_count = 0
    
    for category, templates in question_templates.items():
        for template in templates:
            if question_count >= 8:  # Generate up to 8 sample questions
                break
                
            sample_questions.append({
                'question_text': template,
                'answers': [
                    {'text': f"Comprehensive understanding of {category} aspects", 'is_correct': True},
                    {'text': f"Partial knowledge of {category} elements", 'is_correct': False},
                    {'text': f"Minimal awareness of {category} concepts", 'is_correct': False},
                    {'text': f"No understanding of {category} principles", 'is_correct': False}
                ],
                'explanation': f"This {category} question tests understanding of core concepts in {material_title}.",
                'difficulty': difficulty,
                'question_type': 'MULTIPLE_CHOICE'
            })
            question_count += 1
    
    return sample_questions

def generate_questions(study_material_id):
    """
    Enhanced question generation with multiple types and better content analysis.
    """
    try:
        logger.info(f"Starting enhanced question generation for study material ID: {study_material_id}")
        
        material = StudyMaterial.objects.get(id=study_material_id)
        logger.info(f"Found material: {material.title}")
        
        # Extract text from the file
        file_path = material.file.path
        text = extract_text_from_file(file_path, material.file_type)
        
        if not text or len(text.strip()) < 100:
            logger.error(f"Insufficient text extracted from study material ID: {study_material_id}. Text length: {len(text) if text else 0}")
            # Use diverse sample questions as fallback
            questions = generate_diverse_questions(material.title)
        else:
            # Analyze content complexity
            complexity = analyze_content_complexity(text)
            logger.info(f"Content complexity determined: {complexity}")
            
            # Determine the subject/domain
            subject = detect_subject(material.title, text[:1000])
            logger.info(f"Detected subject: {subject}")
            
            # Split text into manageable chunks
            text_chunks = chunk_text(text, max_chunk_size=3500)
            logger.info(f"Split content into {len(text_chunks)} chunks")
            
            all_questions = []
            
            for i, chunk in enumerate(text_chunks[:3]):  # Process up to 3 chunks
                logger.info(f"Processing chunk {i+1}/{min(len(text_chunks), 3)}")
                
                try:
                    # Create enhanced prompt
                    prompt = create_enhanced_prompt(subject, chunk, material.file_type, complexity, len(text_chunks), i+1)
                    
                    # Generate questions for this chunk
                    response = call_openai_api([
                        {"role": "system", "content": get_system_prompt(subject)},
                        {"role": "user", "content": prompt}
                    ], max_tokens=2500)
                    
                    generated_content = response['choices'][0]['message']['content']
                    logger.info(f"Received response from OpenAI for chunk {i+1}: {len(generated_content)} characters")
                    
                    chunk_questions = parse_enhanced_content(generated_content)
                    logger.info(f"Parsed {len(chunk_questions)} questions from chunk {i+1}")
                    
                    all_questions.extend(chunk_questions)
                    
                except Exception as e:
                    logger.error(f"Error processing chunk {i+1}: {e}")
                    continue
            
            # If no questions generated from chunks, use sample questions
            if not all_questions:
                logger.warning("No questions generated from chunks, using diverse sample questions")
                all_questions = generate_diverse_questions(material.title, subject, complexity)
            
            questions = all_questions[:15]  # Limit to 15 questions maximum
            logger.info(f"Final question set: {len(questions)} questions")
        
        # Save questions and answers to the database
        saved_count = 0
        for i, q in enumerate(questions):
            try:
                difficulty = q.get('difficulty', 'MEDIUM')
                question_type = q.get('question_type', 'MULTIPLE_CHOICE')
                
                question = Question.objects.create(
                    study_material=material,
                    question_text=q['question_text'],
                    question_type=question_type,
                    difficulty=difficulty,
                    explanation=q.get('explanation', '')
                )
                
                for answer in q['answers']:
                    Answer.objects.create(
                        question=question,
                        answer_text=answer['text'],
                        is_correct=answer['is_correct']
                    )
                saved_count += 1
                logger.info(f"Saved question {i+1}: {question.question_text[:50]}...")
                
            except Exception as e:
                logger.error(f"Error saving question {i+1}: {e}")
                continue
        
        if saved_count > 0:
            material.processed = True
            material.save()
            logger.info(f"Successfully generated and saved {saved_count} questions")
            return True
        else:
            logger.error("No questions were saved to the database")
            return False
    
    except StudyMaterial.DoesNotExist:
        logger.error(f"Study material with ID {study_material_id} not found")
        return False
    except Exception as e:
        logger.error(f"Unexpected error generating questions: {e}")
        return False

def detect_subject(title, text_preview):
    """
    Enhanced subject detection with more categories and better scoring.
    """
    subjects = {
        'mathematics': ['math', 'algebra', 'calculus', 'geometry', 'statistics', 'probability', 'equation', 'formula', 'theorem', 'proof', 'derivative', 'integral'],
        'physics': ['physics', 'mechanics', 'electricity', 'magnetism', 'relativity', 'quantum', 'force', 'energy', 'motion', 'wave', 'thermodynamics'],
        'chemistry': ['chemistry', 'organic', 'inorganic', 'biochemistry', 'elements', 'compound', 'molecule', 'reaction', 'bond', 'periodic', 'catalyst'],
        'biology': ['biology', 'anatomy', 'physiology', 'genetics', 'ecology', 'cell', 'organism', 'species', 'evolution', 'dna', 'protein'],
        'computer_science': ['computer', 'programming', 'algorithm', 'data structure', 'software', 'database', 'code', 'function', 'variable', 'api', 'framework'],
        'history': ['history', 'century', 'war', 'civilization', 'empire', 'revolution', 'ancient', 'medieval', 'modern', 'historical', 'timeline'],
        'literature': ['literature', 'novel', 'poetry', 'fiction', 'author', 'character', 'plot', 'theme', 'narrative', 'literary', 'analysis'],
        'economics': ['economics', 'market', 'inflation', 'demand', 'supply', 'fiscal', 'monetary', 'gdp', 'trade', 'investment', 'finance'],
        'psychology': ['psychology', 'behavior', 'cognitive', 'mental', 'brain', 'emotion', 'perception', 'learning', 'memory', 'development'],
        'engineering': ['engineering', 'design', 'construction', 'mechanical', 'electrical', 'civil', 'structural', 'system', 'process', 'technical']
    }
    
    combined_text = (title + " " + text_preview).lower()
    
    # Count occurrences with weighted scoring
    subject_scores = {}
    for subject, keywords in subjects.items():
        score = 0
        for keyword in keywords:
            # Weight title matches more heavily
            title_matches = title.lower().count(keyword) * 3
            text_matches = text_preview.lower().count(keyword)
            score += title_matches + text_matches
        subject_scores[subject] = score
    
    # Return the subject with the highest score
    max_subject = max(subject_scores.items(), key=lambda x: x[1]) if subject_scores else ('general', 0)
    return max_subject[0] if max_subject[1] > 0 else 'general'

def get_system_prompt(subject):
    """
    Get enhanced system prompt based on subject.
    """
    base_prompt = """You are an expert educational assessment creator specializing in generating high-quality, diverse questions for academic materials. Your questions should:

1. Test deep understanding, not just memorization
2. Be clearly written and unambiguous
3. Include various cognitive levels (knowledge, comprehension, application, analysis)
4. Have realistic and plausible distractors
5. Be appropriate for the subject matter and difficulty level"""

    subject_specific = {
        'mathematics': "Focus on problem-solving, conceptual understanding, and application of mathematical principles.",
        'computer_science': "Emphasize algorithmic thinking, code analysis, and practical programming concepts.",
        'science': "Include scientific reasoning, experimental design, and real-world applications.",
        'history': "Focus on cause-and-effect relationships, historical context, and critical analysis.",
        'literature': "Emphasize literary analysis, themes, character development, and author's techniques."
    }
    
    return base_prompt + "\n\n" + subject_specific.get(subject, "Apply general academic assessment principles.")

def create_enhanced_prompt(subject, text, file_type, complexity, total_chunks, chunk_number):
    """
    Create sophisticated prompts for better question generation.
    """
    
    questions_per_chunk = 6 if total_chunks == 1 else 4
    
    prompt = f"""Based on the following {subject} content from chunk {chunk_number} of {total_chunks}, generate {questions_per_chunk} diverse, high-quality assessment questions.

CONTENT COMPLEXITY: {complexity}
CONTENT TYPE: {file_type}

QUESTION REQUIREMENTS:
1. Generate questions of varying cognitive levels:
   - 40% Comprehension/Analysis (understanding concepts, relationships)
   - 30% Application (using knowledge in new situations)
   - 20% Knowledge (factual recall of key information)
   - 10% Synthesis/Evaluation (combining ideas, making judgments)

2. Question types to include:
   - Multiple choice (4 options each)
   - Focus on different aspects: definitions, applications, comparisons, problem-solving

3. Difficulty distribution:
   - {complexity} level as primary difficulty
   - Include 1-2 questions one level easier and 1-2 questions one level harder

4. For each question provide:
   - Clear, specific question text
   - Four plausible answer options (A, B, C, D)
   - Mark the correct answer with (correct)
   - Brief explanation of why the answer is correct
   - Difficulty level (EASY, MEDIUM, HARD)

FORMAT EACH QUESTION EXACTLY LIKE THIS:
Q: [Clear, specific question text]
A: [First option]
B: [Second option]
C: [Third option] (correct)
D: [Fourth option]
Explanation: [Why the correct answer is right and others are wrong]
Difficulty: [EASY/MEDIUM/HARD]

CONTENT TO ANALYZE:
{text[:3500]}

Generate questions that test understanding of the most important concepts, relationships, and applications from this content."""

    # Add subject-specific guidance
    if subject == 'mathematics':
        prompt += "\n\nFocus on: problem-solving steps, formula applications, conceptual understanding, and mathematical reasoning."
    elif subject == 'computer_science':
        prompt += "\n\nFocus on: algorithm logic, code functionality, data structures, and programming concepts."
    elif subject in ['physics', 'chemistry', 'biology']:
        prompt += "\n\nFocus on: scientific principles, experimental applications, cause-and-effect relationships, and real-world examples."
    elif subject == 'history':
        prompt += "\n\nFocus on: chronological understanding, cause-and-effect relationships, historical significance, and contextual analysis."
    elif subject == 'economics':
        prompt += "\n\nFocus on: economic principles, market dynamics, policy implications, and real-world applications."
    
    return prompt

def parse_enhanced_content(content):
    """
    Enhanced parsing with better error handling and question validation.
    """
    logger.info("Parsing enhanced generated content...")
    questions = []
    lines = content.split('\n')
    current_question = None
    
    for line_num, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        logger.debug(f"Processing line {line_num}: {line}")
        
        # New question
        if line.startswith("Q:") or re.match(r"^\d+\.", line):
            # Save previous question if complete
            if current_question and len(current_question.get('answers', [])) >= 4:
                questions.append(current_question)
                logger.debug(f"Added complete question: {current_question['question_text'][:30]}...")
            
            # Start new question
            question_text = line[2:].strip() if line.startswith("Q:") else re.sub(r"^\d+\.\s*", "", line)
            current_question = {
                'question_text': question_text,
                'answers': [],
                'explanation': '',
                'difficulty': 'MEDIUM',
                'question_type': 'MULTIPLE_CHOICE'
            }
        
        # Answer options
        elif current_question and re.match(r"^[A-D][:)]", line):
            option_letter = line[0]
            answer_text = line[2:].strip()
            is_correct = "(correct)" in answer_text.lower()
            
            # Clean up answer text
            answer_text = re.sub(r"\(correct\)", "", answer_text, flags=re.IGNORECASE).strip()
            
            current_question['answers'].append({
                'text': answer_text,
                'is_correct': is_correct
            })
        
        # Explanation
        elif current_question and line.startswith("Explanation:"):
            explanation = line[12:].strip()
            current_question['explanation'] = explanation
        
        # Difficulty
        elif current_question and line.startswith("Difficulty:"):
            difficulty = line[11:].strip().upper()
            if difficulty in ['EASY', 'MEDIUM', 'HARD']:
                current_question['difficulty'] = difficulty
    
    # Add final question
    if current_question and len(current_question.get('answers', [])) >= 4:
        questions.append(current_question)
        logger.debug(f"Added final question: {current_question['question_text'][:30]}...")
    
    # Validate questions
    valid_questions = []
    for q in questions:
        if (len(q.get('answers', [])) == 4 and 
            any(a.get('is_correct') for a in q['answers']) and
            len(q.get('question_text', '')) > 10):
            valid_questions.append(q)
        else:
            logger.warning(f"Skipping invalid question: {q.get('question_text', 'Unknown')[:30]}...")
    
    logger.info(f"Successfully parsed {len(valid_questions)} valid questions out of {len(questions)} total")
    return valid_questions 